from fileinput import filename
from tkinter import*
import tkinter.messagebox 
import os
import random
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
import time
from datetime import date
from xml.dom.minidom import Document
import pyqrcode
from pyqrcode import QRCode
import docx
from docx.shared import Inches
from docx import Document

import firebase_admin
from firebase_admin import credentials
from firebase_admin import db

cred = credentials.Certificate("serviceAccountkey.json")
firebase_admin.initialize_app(cred, {
    'databaseURL' : 'https://restaurant-management-sy-c5c9a-default-rtdb.firebaseio.com/'
})

root=Tk()
root.geometry("1920x1080")
root.title("Restaurant Management System")
root.configure(background="gray")

text_Input = StringVar()
operator = ""

Tops=Frame(root, width=1600,relief=SUNKEN,bg="gray")
Tops.pack(side=TOP)

f1=Frame(root,width=800,height=700,relief=SUNKEN,bg="gray")
f1.pack(side=LEFT)

f2 = Frame(root,width=300, height=700,bg="gray", relief=SUNKEN)
f2.pack(side=RIGHT)

#========================================================================
#                  TIME AND HEADING NAME
#========================================================================

localtime=time.asctime(time.localtime(time.time()))

lblInfo=Label(Tops,font=('algerian',50,'bold'),text="MIET DELICATESSAN ",fg="black",bd=16,bg="gray",anchor='w')
lblInfo.grid(row=0,column=0)
lblInfo=Label(Tops,font=('algerian',20,'bold'),text=localtime,fg="black",bd=16,bg="gray",anchor='w')
lblInfo.grid(row=1,column=0)

#========================================================================
#                  PROGRAM
#========================================================================


def Ref():
    x=random.randint(10908,500876)
    randomRef=str(x)
    rand.set(randomRef)

    if (Idly.get()==""):
        CoIdly=0
    else:
        CoIdly=float(Idly.get())

   
    if (Dosa.get()==""):
        CoDosa=0
    else:
        CoDosa=float(Dosa.get())


    if (IceCream.get()==""):
        CoIceCream=0
    else:
        CoIceCream=float(IceCream.get())


    if (Pulav.get()==""):
        CoPulav=0
    else:
        CoPulav=float(Pulav.get())

        
    if (Tea.get()==""):
        CoTea=0
    else:
        CoTea=float(Tea.get())

     
    if (Drinks.get()==""):
        CoD=0
    else:
        CoD=float(Drinks.get())

                   
    CostofIdly = CoIdly * 25
    CostofDrinks= CoD * 20
    CostofDosa = CoDosa* 25
    CostofIceCream = CoIceCream * 30
    CostPulav = CoPulav* 50
    CostTea = CoTea * 5


    Central_GST= (((CostofIdly+CostofDrinks+CostofDosa+CostofIceCream+CostPulav+CostTea)* 2.5)/100)

    State_GST =(((CostofIdly+CostofDrinks+CostofDosa+CostofIceCream+CostPulav+CostTea)* 2.5)/100)

    Total_cost = (CostofIdly+CostofDrinks+CostofDosa+CostofIceCream+CostPulav+CostTea)

    CostofMeal= "Rs " + str('%.2f' % (CostofIdly+CostofDrinks+CostofDosa+CostofIceCream+CostPulav+CostTea))
    C_gst ="Rs " + str ('%.2f' % Central_GST)
    S_gst = "Rs " + str ('%.2f' % State_GST)
    OverAllCost ="Rs " + str ('%.2f' % (Total_cost+Central_GST+State_GST))


    Sgst.set(S_gst)
    Cost.set(CostofMeal)
    Cgst.set(C_gst)
    Total.set(OverAllCost)
    
def qExit():
    exit=tkinter.messagebox.askyesno("Restaurant Management System","Do you want to exit the system?")
    if exit>0:
       root.destroy()
       return

def Reset():
    Tea.set("")
    Idly.set("")
    Dosa.set("")
    IceCream.set("")
    Pulav.set("")
    Drinks.set("")
    rand.set("")
    Total.set("")
    Sgst.set("")
    Cgst.set("")
    Cost.set("")
    email.set("")
    txtpayslip.delete("1.0",END)

def enterinfo():
  txtpayslip.delete("1.0",END)
  txtpayslip.insert(END,"\t\tPay Slip\n\n")
  if (rand.get()==""):
        random=0
  else:
        txtpayslip.insert(END,"Order No :\t\t"+rand.get()+"\n\n")
  if (Tea.get()==""):
        CoTea=0
  else:
        txtpayslip.insert(END,"Tea :\t\t"+Tea.get()+"\n\n")
  if (Idly.get()==""):
        CoIdly=0
  else:
        txtpayslip.insert(END,"Idli :\t\t"+Idly.get()+"\n\n")
  if (IceCream.get()==""):
        CoIceCream=0
  else:
        txtpayslip.insert(END,"Ice-Cream :\t\t"+IceCream.get()+"\n\n")
  if (Pulav.get()==""):
        CoPulav=0
  else:
        txtpayslip.insert(END,"Pulav :\t\t"+Pulav.get()+"\n\n")
  if (Drinks.get()==""):
        CoD=0
  else:
        txtpayslip.insert(END,"Drinks :\t\t"+Drinks.get()+"\n\n")
  if (Dosa.get()==""):
        CoDosa=0
  else:
        txtpayslip.insert(END,"Dosa :\t\t"+Dosa.get()+"\n\n")
  if (Sgst.get()==""):
        s_gst=0
  else:
        txtpayslip.insert(END,"Sgst :\t\t"+Sgst.get()+"\n\n")
  if (Cgst.get()==""):
        C_gst=0
  else:
        txtpayslip.insert(END,"Cgst :\t\t"+Cgst.get()+"\n\n")
  if (Cost.get()==""):
        costoforder=0
  else:
        txtpayslip.insert(END,"Total cost of the order :\t\t"+Cost.get()+"\n\n") 
  if (Total.get()==""):
        totalorder=0
  else:
        txtpayslip.insert(END,"Total amount to be paid :\t\t"+Total.get()+"\n\n") 


def qprint():
    doc = docx.Document()
    if(rand.get()==""):
        random=0
    else:
        doc.add_paragraph("Order No :\t"+str(rand.get())+"\n")
    if(Tea.get()==""):
        CoTea=0
    else:
        doc.add_paragraph("Tea :\t"+str(Tea.get())+"\n")
    if(Idly.get()==""):
        CoIdly=0
    else:
        doc.add_paragraph("Idli:\t"+str(Idly.get())+"\n")
    if(IceCream.get()==""):
        CoIceCream=0
    else:
        doc.add_paragraph("Ice-Cream :\t"+str(IceCream.get())+"\n")
    if(Pulav.get()==""):
        CoPulav=0
    else:
       doc.add_paragraph("Pulav :\t"+str(Pulav.get())+"\n")
    if(Drinks.get()==""):
        CoD=0
    else:
        doc.add_paragraph("Drinks :\t"+str(Drinks.get())+"\n")
    if(Dosa.get()==""):
        CoDosa=0
    else:
        doc.add_paragraph("Dosa :\t"+str(Dosa.get())+"\n")
    if(Sgst.get()==""):
        s_gst=0
    else:
        doc.add_paragraph("Sgst :\t"+Sgst.get()+"\n")
    if(Cgst.get()==""):
        C_gst=0
    else:
        doc.add_paragraph("Cgst :\t"+Cgst.get()+"\n")
    if(Cost.get()==""):
        costoforder=0
    else:
        doc.add_paragraph("Total cost of the order :\t"+Cost.get()+"\n")
    if(Total.get()==""):
        totalorder=0
    else:
        doc.add_paragraph("Total amount to be paid :\t"+Total.get()+"\n\n\n")
    doc.save('payslip.doc')

def fprint():
    s = "upi://pay?pa=surferab368@oksbi&pn=AbhaySharma&aid=uGICAgIDdt4vKdQ"
    url = pyqrcode.create(s)
    url.svg("myqr.svg", scale = 8)
    url.png('myqr.png', scale = 4)
    doc=docx.Document('payslip.doc')
    doc.add_heading('UPI:')
    doc.add_picture('myqr.png', width=Inches(2), height=Inches(2))
    doc.save('payslip.doc')

    filename='payslip.doc'
    os.startfile(filename, "print")


def save():
    ref = db.reference('info/')
    users_ref = ref.child('payslip')
    users_ref.push({"Total cost of the order" : Cost.get(),
                    "Sgst ":Sgst.get(),
                    "Cgst ":Cgst.get(), 
                    "Total amount paid" : Total.get(),
                    "Order No": rand.get(),
                    "Date of order": str(date.today()),
                    "Time of order": str(time.asctime(time.localtime(time.time())))
                  })

def esend():
    if (email.get() == ""):
        print("no email given")
    else:
        fromaddr = "2020a1r053@mietjammu.in"
        toaddr = email.get()
   
# instance of MIMEMultipart
        msg = MIMEMultipart()
  
# storing the senders email address  
        msg['From'] = fromaddr
  
# storing the receivers email address 
        msg['To'] = toaddr
  
# storing the subject 
        msg['Subject'] = "order info"
  
# string to store the body of the mail
        body = "Here are your order details"
  
# attach the body with the msg instance
        msg.attach(MIMEText(body, 'plain'))
  
# open the file to be sent 
        filename = 'payslip.doc'
        attachment = open(filename, "rb")
  
# instance of MIMEBase and named as p
        p = MIMEBase('application', 'octet-stream')
  
# To change the payload into encoded form
        p.set_payload((attachment).read())
  
# encode into base64
        encoders.encode_base64(p)
   
        p.add_header('Content-Disposition', "attachment; filename= %s" % filename)
  
# attach the instance 'p' to instance 'msg'
        msg.attach(p)
  
# creates SMTP session
        s = smtplib.SMTP('smtp.gmail.com', 587)
  
# start TLS for security
        s.starttls()
  
# Authentication
        s.login(fromaddr, "trickygenius369")
  
# Converts the Multipart msg into a string
        text = msg.as_string()
  
# sending the mail
        s.sendmail(fromaddr, toaddr, text)
  
# terminating the session
        s.quit()
 
#========================================================================
#                  RESTAURANT MENU
#========================================================================

Tea=StringVar()
Idly=StringVar()
Dosa=StringVar()
IceCream=StringVar()
Pulav=StringVar()
Drinks=StringVar()
rand = StringVar()
Cost=StringVar()
Sgst=StringVar()
Cgst=StringVar()
Total=StringVar()
email=StringVar()
DateOfOrder=StringVar()

DateOfOrder.set(time.strftime("%d/%m/%Y"))


lblTea= Label(f1, font=('arail', 16, 'bold'),text="Tea /  Rs5",bd=16,anchor="w",bg="gray")
lblTea.grid(row=0, column=0)
lblTea=Entry(f1, font=('arail',16,'bold'),textvariable=Tea,bd=10,insertwidth=4,bg="white",justify='right')
lblTea.grid(row=0,column=1)

lblDrinks= Label(f1, font=('arail', 16, 'bold'),text="Drinks / Rs20",bd=16,anchor="w",bg="gray")
lblDrinks.grid(row=1, column=0)
txtDrinks=Entry(f1, font=('arail',16,'bold'),textvariable=Drinks,bd=10,insertwidth=4,bg="white",justify='right')
txtDrinks.grid(row=1,column=1)

lblIceCream= Label(f1, font=('arail', 16, 'bold'),text="Ice-Cream / Rs30",bd=16,anchor="w",bg="gray")
lblIceCream.grid(row=2, column=0)
lblIceCream=Entry(f1, font=('arail',16,'bold'),textvariable=IceCream,bd=10,insertwidth=4,bg="white",justify='right')
lblIceCream.grid(row=2,column=1)

lblIdly= Label(f1, font=('arail', 16, 'bold'),text="Idli / Rs25",bd=16,anchor="w",bg="gray")
lblIdly.grid(row=3, column=0)
txtIdly=Entry(f1, font=('arail',16,'bold'),textvariable=Idly,bd=10,insertwidth=4,bg="white",justify='right')
txtIdly.grid(row=3,column=1)

lblDosa= Label(f1, font=('arail', 16, 'bold'),text="Dosa / Rs25",bd=16,anchor="w",bg="gray")
lblDosa.grid(row=4, column=0)
txtDosa=Entry(f1, font=('arail',16,'bold'),textvariable=Dosa,bd=10,insertwidth=4,bg="white",justify='right')
txtDosa.grid(row=4,column=1)

lblPulav= Label(f1, font=('arail', 16, 'bold'),text="Rice-Plate / Rs50",bd=16,anchor="w",bg="gray")
lblPulav.grid(row=5, column=0)
txtPulav=Entry(f1, font=('arail',16,'bold'),textvariable=Pulav,bd=10,insertwidth=4,bg="white",justify='right')
txtPulav.grid(row=5,column=1)


#========================================================================
#                  RESTAURANT BILL INFO
#========================================================================

lblReference= Label(f1, font=('arail', 16, 'bold'),text="Order No",bd=16,anchor="w",bg="gray")
lblReference.grid(row=0, column=2)
txtReference=Entry(f1, font=('arail',16,'bold'),textvariable=rand,bd=10,insertwidth=4,bg="white",justify='right')
txtReference.grid(row=0,column=3)

lblCost= Label(f1, font=('arail', 16, 'bold'),text="Cost of Meal",bd=16,anchor="w",bg="gray")
lblCost.grid(row=1, column=2)
txtCost=Entry(f1, font=('arail',16,'bold'),textvariable=Cost,bd=10,insertwidth=4,bg="white",justify='right')
txtCost.grid(row=1,column=3)


lblSgst= Label(f1, font=('arail', 16, 'bold'),text="SGST",bd=16,anchor="w",bg="gray")
lblSgst.grid(row=2, column=2)
txtSgst=Entry(f1, font=('arail',16,'bold'),textvariable=Sgst,bd=10,insertwidth=4,bg="white",justify='right')
txtSgst.grid(row=2,column=3)


lblCgst= Label(f1, font=('arail', 16, 'bold'),text="CGST",bd=16,anchor="w",bg="gray")
lblCgst.grid(row=3, column=2)
txtCgst=Entry(f1, font=('arail',16,'bold'),textvariable=Cgst,bd=10,insertwidth=4,bg="white",justify='right')
txtCgst.grid(row=3,column=3)

lblTotalCost= Label(f1, font=('arail', 16, 'bold'),text="Total Cost",bd=16,anchor="w",bg="gray")
lblTotalCost.grid(row=4, column=2)
txtTotalCost=Entry(f1, font=('arail',16,'bold'),textvariable=Total,bd=10,insertwidth=4,bg="white",justify='right')
txtTotalCost.grid(row=4,column=3)

lblemail= Label(f1, font=('arail', 16, 'bold'),text="E-mail",bd=16,anchor="w",bg="gray")
lblemail.grid(row=5, column=2)
txtemail=Entry(f1, font=('arail',16,'bold'),textvariable=email,bd=10,insertwidth=4,bg="white",justify='right')
txtemail.grid(row=5,column=3)

#========================================================================
#                  Text Widget
#========================================================================
payslip=Label(f2,textvariable=DateOfOrder,font=('arial',21,'bold'),fg="black",bg="gray").grid(row=0,column=0)
txtpayslip=Text(f2,height=22,width=34,bd=16,font=('arial',13,'bold'),fg="green",bg="gray")
txtpayslip.grid(row=1,column=0)

#========================================================================
#                  BUTTONS
#========================================================================
btnTotal=Button(f1,padx=16,pady=8,bd=16,fg="black",font=('arail',16,'bold'),width=10,text="Total",bg="gray",command=Ref).grid(row=7,column=1)

btnReset=Button(f1,padx=16,pady=8,bd=16,fg="black",font=('arail',16,'bold'),width=10,text="Reset",bg="gray",command=Reset).grid(row=7,column=2)

btnpayslip=Button(f1,padx=16,pady=8,bd=16,fg="black",font=('arail',16,'bold'),width=10,text="Payslip",bg="gray",command=enterinfo).grid(row=7,column=3)

btnExit=Button(f1,padx=16,pady=8,bd=16,fg="black",font=('arail',16,'bold'),width=10,text="Exit",bg="gray",command=qExit).grid(row=8,column=1)

btnprint=Button(f1,padx=16,pady=8,bd=16,fg="black",font=('arail',16,'bold'),width=10,text="Save",bg="gray",command=save).grid(row=8,column=2)

btnprint=Button(f1,padx=16,pady=8,bd=16,fg="black",font=('arail',16,'bold'),width=10,text="Print",bg="gray",command=lambda:[qprint(), fprint(), esend()]).grid(row=8,column=3)


root.mainloop()



